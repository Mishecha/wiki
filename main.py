import requests
from bs4 import BeautifulSoup
from docx import Document


def parse_wikipedia_article(url):
    try:
        response = requests.get(url)
        response.raise_for_status()  
    except requests.exceptions.RequestException as e:
        print(f"Ошибка при получении статьи: {e}")
        return

    soup = BeautifulSoup(response.content, 'html.parser')
    doc = Document()

    title = soup.find(id="firstHeading").text
    doc.add_heading(title, 1)

    content = soup.find_all(['h2', 'h3', 'p'])
    for element in content:
        if element.name == 'h2':
            doc.add_heading(element.text.strip(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text.strip(), level=3)
        elif element.name == 'p':
            doc.add_paragraph(element.text.strip())
        if element.id == 'См._также':
            break
        if element.id == 'See_also':
            break

    filename = f"{title.text.strip()}.docx"
    doc.save(filename)


def main():
    url = input("Введите ссылку на статью Википедии: ").sprip()
    parse_wikipedia_article(url)


if __name__ == "__main__":
    main()
